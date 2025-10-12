'''
This file takes in word and pdf files, and strips all personal information from the header
'''


# 1. Import all necessary libraries
from pathlib import Path
import docx
from pypdf import PdfReader, PdfWriter
import spacy
import fitz

# 2. Load the spaCy NER model once when the script starts
print("Loading spaCy NLP model...")
nlp = spacy.load("en_core_web_sm")
print("Model loaded successfully.")


def process_docx(input_path, output_path_folder, nlp_model):
    """
    Opens a .docx, intelligently finds and removes paragraphs 
    containing person names from the body and headers, and saves it.
    """
    try:
        document = docx.Document(input_path)
        
        # --- Clean the main body (first 4 paragraphs) ---
        body_paragraphs_to_check = document.paragraphs[:4]
        body_paragraphs_to_delete = []
        for p in body_paragraphs_to_check:
            if p.text.strip():
                doc = nlp_model(p.text)
                for ent in doc.ents:
                    if ent.label_ == "PERSON":
                        body_paragraphs_to_delete.append(p)
                        break 
        for p in body_paragraphs_to_delete:
            p._element.getparent().remove(p._element)

        # --- Clean Headers ---
        header_paragraphs_deleted_count = 0
        for section in document.sections:
            header = section.header
            header_paragraphs_to_delete = []
            for p in header.paragraphs:
                if p.text.strip():
                    doc = nlp_model(p.text)
                    for ent in doc.ents:
                        if ent.label_ == "PERSON":
                            header_paragraphs_to_delete.append(p)
                            break
            for p in header_paragraphs_to_delete:
                p._element.getparent().remove(p._element)
                header_paragraphs_deleted_count += 1

        output_file_path = output_path_folder / input_path.name
        document.save(output_file_path)
        total_removed = len(body_paragraphs_to_delete) + header_paragraphs_deleted_count
        print(f"  -> Saved DOCX: {output_file_path.name} ({total_removed} lines removed)")
    except Exception as e:
        print(f"  -> Error processing {input_path.name}: {e}")


def process_pdf(input_path, output_path_folder):
    """Opens a .pdf, covers the top margin of each page using PyMuPDF, and saves it."""
    try:
        # Open the source PDF with fitz
        doc = fitz.open(input_path)

        # IMPORTANT: Adjust this value as needed.
        TOP_MARGIN_INCHES = 1.25
        margin_to_cover = TOP_MARGIN_INCHES * 72  # 1 inch = 72 points

        # Loop through all pages in the PDF
        for page in doc:
            # The rectangle area to cover (x0, y0, x1, y1)
            # Coordinates start from the top-left corner in PyMuPDF
            rect = fitz.Rect(0, 0, page.rect.width, margin_to_cover)

            # Draw a white rectangle to cover the area
            # 'color' is the border, 'fill' is the inside color. Both are white.
            page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))

        # Construct the full path for the new file
        output_file_path = output_path_folder / input_path.name
        
        # Save the modified PDF, garbage collection cleans up memory
        doc.save(output_file_path, garbage=4, deflate=True, clean=True)
        doc.close()
            
        print(f"  -> Saved processed file: {output_file_path.name}")

    except Exception as e:
        print(f"  -> Error processing {input_path.name}: {e}")
    """Opens a .pdf, covers the top margin of each page, and saves it."""
    try:
        reader = PdfReader(input_path)
        writer = PdfWriter()
        TOP_MARGIN_INCHES = 1.25
        margin_to_cover = TOP_MARGIN_INCHES * 72

        for page in reader.pages:
            page_height = page.mediabox.height
            page_width = page.mediabox.width
            page.draw_rectangle(
                rect=(0, float(page_height) - margin_to_cover, float(page_width), float(page_height)),
                fill_color=(1, 1, 1)
            )
            writer.add_page(page)

        output_file_path = output_path_folder / input_path.name
        with open(output_file_path, "wb") as f:
            writer.write(f)
        print(f"  -> Saved PDF: {output_file_path.name}")
    except Exception as e:
        print(f"  -> Error processing {input_path.name}: {e}")


def main():
    """
    Main function to define folders and process all files.
    """
    input_folder = Path("./originals")
    output_folder = Path("./processed")
    output_folder.mkdir(exist_ok=True)

    print(f"\nScanning for files in '{input_folder}'...")

    for file_path in input_folder.iterdir():
        '''
        if file_path.suffix == ".docx":
            print(f"Processing DOCX file: {file_path.name}")
            process_docx(file_path, output_folder, nlp)
        '''

        if file_path.suffix == ".pdf":
            print(f"Processing PDF file: {file_path.name}")
            process_pdf(file_path, output_folder)
        
        else:
            print(f"Skipping unsupported file: {file_path.name}")
    
    print("\nScript finished.")


# This block ensures the main() function runs only when you execute the script directly
if __name__ == "__main__":
    main()